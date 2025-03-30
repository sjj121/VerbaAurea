import os
from docx import Document
import sys
import nltk
from nltk.tokenize import sent_tokenize
import jieba
from config_manager import load_config, show_config, edit_config

# 下载必要的nltk数据（第一次运行需要）
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')


def is_sentence_boundary(text_before, text_after):
    """判断两段文本之间是否为句子边界"""
    # 检查前文是否以句号结尾
    if text_before.endswith(('。', '！', '？', '.', '!', '?', '；', ';')):
        return True

    # 使用jieba进行更精确的句子边界检测
    combined_text = text_before + " " + text_after
    try:
        # 区分中英文进行句子分割
        if any(u'\u4e00' <= char <= u'\u9fff' for char in combined_text):
            # 中文文本，使用jieba分句
            sentences = list(jieba.cut(combined_text))
            # 检查分词结果中是否有明显的句子边界
            for i, word in enumerate(sentences[:-1]):
                if word in ['。', '！', '？', '.', '!', '?', '；', ';']:
                    # 检查这个边界是否接近text_before和text_after的连接处
                    before_seg = ''.join(sentences[:i+1])
                    if (len(before_seg) - len(text_before)) < 5:
                        return True
        else:
            # 英文文本，使用NLTK
            sentences = sent_tokenize(combined_text)
            # 检查是否能在组合文本中找到明确的句子分界点
            for sentence in sentences:
                if text_before.endswith(sentence) or text_after.startswith(sentence):
                    return True
    except:
        pass

    return False


def find_nearest_sentence_boundary(paragraphs_info, current_index, search_window=5):
    """寻找距离当前位置最近的句子边界"""
    best_index = -1
    min_distance = float('inf')

    # 向前查找
    for i in range(max(0, current_index - search_window), current_index + 1):
        if i > 0 and is_sentence_boundary(paragraphs_info[i - 1]['text'], paragraphs_info[i]['text']):
            distance = current_index - i
            if 0 <= distance < min_distance:
                min_distance = distance
                best_index = i

    # 向后查找
    for i in range(current_index + 1, min(len(paragraphs_info), current_index + search_window + 1)):
        if i > 0 and is_sentence_boundary(paragraphs_info[i - 1]['text'], paragraphs_info[i]['text']):
            distance = i - current_index
            if distance < min_distance:
                min_distance = distance
                best_index = i

    return best_index


def get_paragraph_semantic_block(paragraphs_info, start_idx, end_idx):
    """获取段落范围的语义块文本"""
    block_text = ""
    for i in range(start_idx, end_idx + 1):
        block_text += paragraphs_info[i]['text'] + " "
    return block_text.strip()


def insert_split_markers(input_file, output_file, config):
    """
    在Word文档中根据算法自动嵌入<!--split-->标记

    使用多种策略确定分隔点:
    1. 标题识别
    2. 语义结构分析
    3. 长度控制
    4. 自然段落边界
    5. 句子完整性保证
    """
    # 从配置中提取参数
    doc_settings = config["document_settings"]
    proc_options = config["processing_options"]
    adv_settings = config["advanced_settings"]

    max_length = doc_settings["max_length"]
    min_length = doc_settings["min_length"]
    sentence_integrity_weight = doc_settings["sentence_integrity_weight"]
    debug_mode = proc_options["debug_mode"]
    search_window = adv_settings["search_window"]
    min_split_score = adv_settings["min_split_score"]
    heading_score_bonus = adv_settings["heading_score_bonus"]
    sentence_end_score_bonus = adv_settings["sentence_end_score_bonus"]
    length_score_factor = adv_settings["length_score_factor"]

    print(f"正在处理文档: {input_file}")

    # 如果配置为跳过已存在文件，且输出文件已存在，则跳过处理
    if proc_options["skip_existing"] and os.path.exists(output_file):
        print(f"跳过已存在文件: {output_file}")
        return True

    try:
        doc = Document(input_file)
    except Exception as e:
        print(f"无法打开文档 {input_file}: {str(e)}")
        return False

    # 创建新文档
    new_doc = Document()

    # 第一步: 文档分析阶段
    paragraphs_info = []
    total_text = ""

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:  # 跳过空段落，但记录它们以保持索引一致
            paragraphs_info.append({
                'index': i,
                'text': "",
                'length': 0,
                'is_heading': False,
                'is_list_item': False,
                'ends_with_period': False,
                'length_category': "empty"
            })
            continue

        # 判断段落类型
        is_heading = paragraph.style.name.startswith(('Heading', '标题'))
        is_list_item = text.startswith(('•', '-', '*', '1.', '2.', '3.')) or (
                len(text) > 2 and text[0].isdigit() and text[1] == '.')

        # 段落长度分类
        length_category = "short" if len(text) < 50 else "medium" if len(text) < 200 else "long"

        # 尝试判断段落是否为完整句子结束
        ends_with_period = text.endswith(('。', '！', '？', '.', '!', '?', '；', ';'))

        # 将段落信息保存起来
        paragraphs_info.append({
            'index': i,
            'text': text,
            'length': len(text),
            'is_heading': is_heading,
            'is_list_item': is_list_item,
            'ends_with_period': ends_with_period,
            'length_category': length_category
        })

        total_text += text + " "

    # 使用NLTK尝试进行句子拆分（用于辅助判断语义边界）
    try:
        # 判断文本主要是中文还是英文
        if any(u'\u4e00' <= char <= u'\u9fff' for char in total_text):
            # 中文文本使用jieba分词和分句
            words = list(jieba.cut(total_text))
            sentences = []
            temp = ""
            for word in words:
                temp += word
                if word in ['。', '！', '？', '.', '!', '?']:
                    sentences.append(temp)
                    temp = ""
            if temp:  # 添加最后一个可能不完整的句子
                sentences.append(temp)
        else:
            # 英文文本使用NLTK
            sentences = sent_tokenize(total_text)

        avg_sentence_length = len(total_text) / len(sentences) if sentences else 0
    except:
        avg_sentence_length = 80  # 如果失败，使用默认值

    # 预处理阶段：尝试将连续的短段落组合成语义块
    semantic_blocks = []
    temp_block = {'text': '', 'paragraphs': []}

    for i, para_info in enumerate(paragraphs_info):
        if para_info['length'] == 0:  # 跳过空段落
            continue

        # 如果是标题或列表项开始，结束前一个块
        if para_info['is_heading'] or (
                para_info['is_list_item'] and (i == 0 or not paragraphs_info[i - 1]['is_list_item'])):
            if temp_block['text']:
                semantic_blocks.append(temp_block)
                temp_block = {'text': '', 'paragraphs': []}

        # 添加到当前块
        temp_block['text'] += para_info['text'] + ' '
        temp_block['paragraphs'].append(i)

        # 如果段落结束一个完整句子，且不是短段落，也考虑结束当前块
        if para_info['ends_with_period'] and para_info['length'] > 100:
            semantic_blocks.append(temp_block)
            temp_block = {'text': '', 'paragraphs': []}

    # 添加最后一个块
    if temp_block['text']:
        semantic_blocks.append(temp_block)

    if debug_mode:
        print(f"文档共有 {len(paragraphs_info)} 个段落，组合成 {len(semantic_blocks)} 个语义块")

    # 第二步: 确定分隔点
    split_points = []
    current_length = 0
    last_potential_split = -1

    for i, para_info in enumerate(paragraphs_info):
        if para_info['length'] == 0:  # 跳过空段落
            continue

        current_length += para_info['length']

        # 潜在的分隔点评分系统 (0-10分，越高越适合作为分隔点)
        split_score = 0

        # 1. 标题是最佳分隔点
        if para_info['is_heading']:
            split_score += heading_score_bonus

        # 2. 句子完整性检查
        if para_info['ends_with_period']:
            split_score += sentence_end_score_bonus

        # 3. 确保分割点在句子边界
        if i > 0 and is_sentence_boundary(paragraphs_info[i - 1]['text'], para_info['text']):
            split_score += 8 * sentence_integrity_weight / 8.0  # 使用传入的权重调整
        else:
            split_score -= 10  # 严重惩罚非句子边界的分割

        # 4. 长度已达到理想范围加分
        if current_length >= min_length:
            split_score += min(4, (current_length - min_length) // length_score_factor)
        elif current_length < min_length * 0.7:  # 如果太短则减分
            split_score -= 5

        # 5. 列表项开始或结束是好的分隔点
        if i > 0 and ((para_info['is_list_item'] and not paragraphs_info[i - 1]['is_list_item']) or
                      (not para_info['is_list_item'] and paragraphs_info[i - 1]['is_list_item'] and
                       paragraphs_info[i - 1]['ends_with_period'])):
            split_score += 3

        # 6. 避免相邻分隔点
        if split_points and i - split_points[-1] < 3:
            split_score -= 8

        # 7. 如果长度超过最大值，增加分割倾向
        if current_length > max_length:
            split_score += 4

        if debug_mode and split_score >= 0:
            print(f"段落 {i}: 得分={split_score:.1f}, 文本预览: '{para_info['text'][:50]}...'")

        # 记录潜在的分隔点
        if split_score >= min_split_score and i > 0:  # 至少要设定分以上才考虑作为分隔点
            if debug_mode:
                print(f"选择分割点: {i}, 得分: {split_score:.1f}")

            split_points.append(i)
            current_length = para_info['length']  # 重置长度计数
            last_potential_split = i
        elif current_length > max_length * 1.5:
            # 长度已经超过最大限制的1.5倍，需要找一个合适的分割点
            best_boundary_index = find_nearest_sentence_boundary(paragraphs_info, i, search_window)

            if best_boundary_index >= 0 and (not split_points or best_boundary_index > split_points[-1]):
                if debug_mode:
                    print(f"超长分段，选择最近的句子边界: {best_boundary_index}")

                split_points.append(best_boundary_index)
                # 重新计算当前长度
                if best_boundary_index < i:
                    current_length = sum(p['length'] for p in paragraphs_info[best_boundary_index:i + 1])
                else:
                    current_length = para_info['length']
                last_potential_split = best_boundary_index
            elif i - last_potential_split > 3:
                # 实在找不到合适的句子边界，只能在当前位置分割
                if debug_mode:
                    print(f"警告：在段落 {i} 处强制分割，未找到合适的句子边界")

                split_points.append(i)
                current_length = para_info['length']
                last_potential_split = i

    # 后处理：检查所有分割点确保不会打断句子
    final_split_points = []
    for split_point in split_points:
        # 获取分割点前后的文本
        before_text = paragraphs_info[split_point - 1]['text'] if split_point > 0 else ""
        after_text = paragraphs_info[split_point]['text']

        # 检查是否为句子边界
        if is_sentence_boundary(before_text, after_text):
            final_split_points.append(split_point)
        else:
            # 寻找附近更合适的分割点
            best_point = find_nearest_sentence_boundary(paragraphs_info, split_point, search_window)
            if best_point >= 0 and best_point not in final_split_points:
                if debug_mode:
                    print(f"修正分割点: {split_point} -> {best_point}")
                final_split_points.append(best_point)
            else:
                # 无法找到更好的点，保留原分割点但记录警告
                if debug_mode:
                    print(f"警告: 无法找到比 {split_point} 更好的句子边界")
                final_split_points.append(split_point)

    # 确保分割点有序
    final_split_points = sorted(list(set(final_split_points)))

    if debug_mode:
        print(f"最终分割点: {final_split_points}")

    # 第三步: 创建带有分隔符的新文档
    split_marker_count = 0
    current_para_index = 0

    for i, paragraph in enumerate(doc.paragraphs):
        # 检查是否是分隔点
        if i in final_split_points:
            # 添加分隔符
            new_para = new_doc.add_paragraph("<!--split-->")
            split_marker_count += 1

        # 添加当前段落
        text = paragraph.text
        new_para = new_doc.add_paragraph(text)

        # 复制格式
        try:
            if paragraph.style:
                new_para.style = paragraph.style
            new_para.alignment = paragraph.alignment

            # 复制段落内的文本格式
            for j in range(min(len(paragraph.runs), len(new_para.runs))):
                src_run = paragraph.runs[j]
                dst_run = new_para.runs[j]

                for attr in ['bold', 'italic', 'underline']:
                    if hasattr(src_run, attr):
                        setattr(dst_run, attr, getattr(src_run, attr))

                # 复制字体属性
                if hasattr(src_run, 'font') and hasattr(dst_run, 'font'):
                    for attr in ['size', 'name', 'color']:
                        try:
                            if hasattr(src_run.font, attr):
                                src_value = getattr(src_run.font, attr)
                                if src_value:
                                    setattr(dst_run.font, attr, src_value)
                        except:
                            pass
        except Exception as e:
            if debug_mode:
                print(f"  警告: 复制格式时出错: {str(e)}")

    # 处理表格
    try:
        for table in doc.tables:
            try:
                rows = len(table.rows)
                cols = 0
                if rows > 0:
                    cols = len(table.rows[0].cells)

                if rows > 0 and cols > 0:
                    new_table = new_doc.add_table(rows=rows, cols=cols)

                    try:
                        new_table.style = table.style
                    except:
                        pass

                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                                try:
                                    new_cell = new_table.rows[i].cells[j]
                                    if cell.text:
                                        new_cell.text = cell.text
                                except:
                                    pass
            except Exception as e:
                print(f"  警告: 处理表格时出错: {str(e)}")
    except Exception as e:
        print(f"  警告: 处理文档中的表格时出错: {str(e)}")

    # 保存文档
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    try:
        new_doc.save(output_file)
        print(f"已保存: {output_file}，插入了 {split_marker_count} 个分隔符")
        return True
    except Exception as e:
        print(f"保存文档 {output_file} 时出错: {str(e)}")
        return False


def process_all_documents(config):
    """处理当前目录下所有Word文档"""
    # 获取脚本当前路径
    current_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
    if not current_dir:  # 如果为空，使用当前工作目录
        current_dir = os.getcwd()

    # 创建输出目录
    output_folder = config["processing_options"]["output_folder"]
    output_base_dir = os.path.join(current_dir, output_folder)
    os.makedirs(output_base_dir, exist_ok=True)

    total_files = 0
    processed_files = 0
    failed_files = []

    # 遍历当前目录及子目录
    for root, dirs, files in os.walk(current_dir):
        # 跳过输出文件夹
        if output_folder in root:
            continue

        # 创建相对路径
        rel_path = os.path.relpath(root, current_dir)
        if rel_path == ".":  # 当前目录
            rel_path = ""

        # 处理当前目录下的所有Word文档
        for file in files:
            if file.endswith(('.docx', '.doc')) and not file.startswith('~$'):  # 排除临时文件
                total_files += 1

                # 构建输入和输出路径
                input_path = os.path.join(root, file)

                # 构建输出路径，保持原始目录结构
                if rel_path:
                    output_dir = os.path.join(output_base_dir, rel_path)
                    os.makedirs(output_dir, exist_ok=True)
                else:
                    output_dir = output_base_dir

                output_path = os.path.join(output_dir, file)

                # 处理文档
                try:
                    if insert_split_markers(input_path, output_path, config):
                        processed_files += 1
                    else:
                        failed_files.append(input_path)
                except Exception as e:
                    print(f"处理 {input_path} 时出错: {str(e)}")
                    failed_files.append(input_path)

    return total_files, processed_files, failed_files


def main():
    """主函数"""
    print("Word文档智能分割工具 v2.0")
    print("============================")
    print("该工具会根据智能算法在Word文档中插入<!--split-->分隔符")
    print("优化版本: 使用配置文件管理参数，增强句子完整性保护，避免在句子中间分割")

    # 检查必要的库
    required_libs = ['nltk', 'jieba']
    missing_libs = []

    for lib in required_libs:
        try:
            __import__(lib)
        except ImportError:
            missing_libs.append(lib)

    if missing_libs:
        print(f"\n警告: 未安装以下库: {', '.join(missing_libs)}")
        print("为获得最佳分割效果，建议安装这些库:")
        for lib in missing_libs:
            print(f"  pip install {lib}")
        print()

    # 加载配置
    config = load_config()

    # 显示主菜单
    while True:
        print("\n=== 主菜单 ===")
        print("1. 开始处理文档")
        print("2. 查看当前配置")
        print("3. 编辑配置")
        print("4. 退出")

        choice = input("\n请选择操作 [1-4]: ").strip()

        if choice == '1':
            # 开始处理文档
            print("\n开始处理文档...")
            total_files, processed_files, failed_files = process_all_documents(config)

            print(
                f"\n处理完成! 共找到 {total_files} 个Word文档，成功处理 {processed_files} 个，失败 {len(failed_files)} 个。")
            print(f"处理后的文档已保存在当前目录下的'{config['processing_options']['output_folder']}'文件夹中。")

            if failed_files:
                print("\n以下文件处理失败:")
                for file in failed_files:
                    print(f" - {file}")

            input("\n按Enter键继续...")

        elif choice == '2':
            # 查看当前配置
            show_config()
            input("\n按Enter键继续...")

        elif choice == '3':
            # 编辑配置
            edit_config()
            # 重新加载配置以确保使用最新的值
            config = load_config()

        elif choice == '4':
            # 退出
            print("\n谢谢使用！")
            break

        else:
            print("\n无效选择，请重试")


if __name__ == "__main__":
    main()