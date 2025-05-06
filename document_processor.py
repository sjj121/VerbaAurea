#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文档处理核心模块
处理Word文档，在适当位置插入分隔标记
"""

from docx import Document
from text_analysis import extract_elements_info
import os
from text_analysis import (
    is_sentence_boundary,
    find_nearest_sentence_boundary
)
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl


def insert_split_markers(input_file, output_file, config):
    """
    在Word文档中根据算法自动嵌入<!--split-->标记

    使用多种策略确定分隔点:
    1. 标题识别
    2. 语义结构分析
    3. 长度控制
    4. 自然段落边界
    5. 句子完整性保证
    """
    # 从配置中提取参数
    doc_settings = config["document_settings"]
    proc_options = config["processing_options"]
    adv_settings = config["advanced_settings"]

    max_length = doc_settings["max_length"]
    min_length = doc_settings["min_length"]
    sentence_integrity_weight = doc_settings["sentence_integrity_weight"]
    debug_mode = proc_options["debug_mode"]
    search_window = adv_settings["search_window"]
    min_split_score = adv_settings["min_split_score"]
    heading_score_bonus = adv_settings["heading_score_bonus"]
    sentence_end_score_bonus = adv_settings["sentence_end_score_bonus"]
    length_score_factor = adv_settings["length_score_factor"]

    if debug_mode:
        print(f"正在处理文档: {input_file}")

    # 如果配置为跳过已存在文件，且输出文件已存在，则跳过处理
    if proc_options["skip_existing"] and os.path.exists(output_file):
        if debug_mode:
            print(f"跳过已存在文件: {output_file}")
        return True

    try:
        doc = Document(input_file)
    except Exception as e:
        print(f"无法打开文档 {input_file}: {str(e)}")
        return False

    # 创建新文档
    new_doc = Document()

    # 统一提取段落 + 表格
    table_factor = config.get("document_settings", {}).get("table_length_factor", 1.0)
    elements_info = extract_elements_info(doc, table_factor, debug_mode)

    if debug_mode:
        print(f"文档共有 {len(elements_info)} 个元素（段落 + 表格）")

    # 确定分隔点
    split_points = find_split_points(
        elements_info,
        max_length, min_length,
        sentence_integrity_weight, search_window,
        min_split_score, heading_score_bonus,
        sentence_end_score_bonus, length_score_factor,
        debug_mode
    )

    # 后处理：检查所有分割点确保不会打断句子
    final_split_points = refine_split_points(
        elements_info, split_points, search_window, debug_mode
    )

    if debug_mode:
        print(f"最终分割点: {final_split_points}")

    # 创建带有分隔符的新文档
    result = create_output_document(
        doc,
        new_doc,
        final_split_points,
        output_file,
        debug_mode
    )

    return result


# ---------- 重新实现分割三大函数 ----------
def find_split_points(elements_info, max_length, min_length,
                      sentence_integrity_weight, search_window,
                      min_split_score, heading_score_bonus,
                      sentence_end_score_bonus, length_score_factor,
                      debug_mode):
    split_points = []
    current_length = 0
    last_potential = -1

    for idx, elem in enumerate(elements_info):
        if elem['length'] == 0:
            continue

        current_length += elem['length']
        score = calculate_split_score(
            idx, elem, elements_info, current_length,
            min_length, max_length, sentence_integrity_weight,
            heading_score_bonus, sentence_end_score_bonus,
            length_score_factor, split_points
        )

        if debug_mode:
            preview = (elem['text'][:30] + '...') if elem['text'] else '[table]'
            print(f"  #{idx:03d} ({elem['type']}) len={elem['length']} score={score:.1f} {preview}")

        if score >= min_split_score and idx > 0:
            split_points.append(idx)
            current_length = 0
            last_potential = idx
        elif current_length > max_length * 1.5:
            best = find_nearest_sentence_boundary(elements_info, idx, search_window)
            if best >= 0 and (not split_points or best > split_points[-1]):
                split_points.append(best)
                current_length = 0
                last_potential = best
            elif idx - last_potential > 3:
                split_points.append(idx)
                current_length = 0
                last_potential = idx

    return split_points


def calculate_split_score(idx, elem, elements_info, current_length,
                          min_length, max_length, sentence_integrity_weight,
                          heading_score_bonus, sentence_end_score_bonus,
                          length_score_factor, split_points):
    score = 0
    if elem['type'] == 'para':
        if elem['is_heading']:
            score += heading_score_bonus
        if elem['ends_with_period']:
            score += sentence_end_score_bonus
        # 句子边界检查（仅段落间）
        if idx > 0 and elements_info[idx-1]['type'] == 'para' and \
           is_sentence_boundary(elements_info[idx-1]['text'], elem['text']):
            score += sentence_integrity_weight
        else:
            score -= 10
    else:
        # 表格：天然边界，可在其前后分段
        score += 6

    # 长度因子
    if current_length >= min_length:
        score += min(4, (current_length - min_length)//length_score_factor)
    elif current_length < min_length*0.7:
        score -= 5

    # 避免过近
    if split_points and idx - split_points[-1] < 3:
        score -= 8
    if current_length > max_length:
        score += 4
    return score


def refine_split_points(elements_info, split_points, search_window, debug_mode):
    refined = []
    for sp in split_points:
        before = elements_info[sp-1]['text'] if sp>0 else ''
        after  = elements_info[sp]['text']
        if elements_info[sp-1]['type']=='para' and elements_info[sp]['type']=='para' and \
           not is_sentence_boundary(before, after):
            best = find_nearest_sentence_boundary(elements_info, sp, search_window)
            refined.append(best if best>=0 else sp)
        else:
            refined.append(sp)
    return sorted(set(refined))



def copy_single_table(table, new_doc, debug_mode):
    """复制单个表格"""
    if table is None:  # ← 兜底
        return
    try:
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if rows > 0 else 0

        if rows > 0 and cols > 0:
            new_table = new_doc.add_table(rows=rows, cols=cols)

            try:
                new_table.style = table.style
            except:
                pass

            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                        try:
                            new_cell = new_table.rows[i].cells[j]
                            if cell.text:
                                new_cell.text = cell.text
                        except:
                            pass
    except Exception as e:
        if debug_mode:
            print(f"  警告: 处理表格时出错: {str(e)}")


def create_output_document(doc, new_doc, split_points, output_file, debug_mode):
    split_marker_cnt = 0
    para_iter = iter(doc.paragraphs)
    tbl_iter  = iter(doc.tables)
    next_para = next(para_iter, None)
    next_tbl  = next(tbl_iter, None)
    idx = -1

    # 将 Word DOM 再次顺序遍历
    for el in doc._element.body:
        idx += 1
        if idx in split_points:
            new_doc.add_paragraph("<!--split-->")
            split_marker_cnt += 1

        if isinstance(el, CT_P):
            # —— 段落 ——
            copy_paragraph(next_para, new_doc, debug_mode)
            next_para = next(para_iter, None)
        elif isinstance(el, CT_Tbl):
            # —— 表格 ——
            copy_single_table(next_tbl, new_doc, debug_mode)
            next_tbl = next(tbl_iter, None)

    # 保存
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    new_doc.save(output_file)
    if debug_mode:
        print(f"✓ 保存: {output_file} (split={split_marker_cnt})")
    return True



def copy_paragraph(src_para, new_doc, debug_mode):
    """复制段落内容和格式"""
    text = src_para.text
    new_para = new_doc.add_paragraph(text)

    # 复制格式
    try:
        if src_para.style:
            new_para.style = src_para.style
        new_para.alignment = src_para.alignment

        # 复制段落内的文本格式
        for j in range(min(len(src_para.runs), len(new_para.runs))):
            src_run = src_para.runs[j]
            dst_run = new_para.runs[j]

            for attr in ['bold', 'italic', 'underline']:
                if hasattr(src_run, attr):
                    setattr(dst_run, attr, getattr(src_run, attr))

            # 复制字体属性
            if hasattr(src_run, 'font') and hasattr(dst_run, 'font'):
                for attr in ['size', 'name', 'color']:
                    try:
                        if hasattr(src_run.font, attr):
                            src_value = getattr(src_run.font, attr)
                            if src_value:
                                setattr(dst_run.font, attr, src_value)
                    except:
                        pass
    except Exception as e:
        if debug_mode:
            print(f"  警告: 复制格式时出错: {str(e)}")

