#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文本分析模块
负责段落分析、语义结构识别、句子边界检测等
"""

import jieba
import nltk
from nltk.tokenize import sent_tokenize
import functools
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
# 初始化 NLTK 和 Jieba
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')


@functools.lru_cache(maxsize=1024)  # 缓存结果以提高性能
def is_sentence_boundary(text_before, text_after):
    """判断两段文本之间是否为句子边界"""
    # 检查前文是否以句号结尾
    if text_before.endswith(('。', '！', '？', '.', '!', '?', '；', ';')):
        return True

    # 使用jieba进行更精确的句子边界检测
    combined_text = text_before + " " + text_after
    try:
        # 区分中英文进行句子分割
        if any(u'\u4e00' <= char <= u'\u9fff' for char in combined_text):
            # 中文文本，使用jieba分句
            sentences = list(jieba.cut(combined_text))
            # 检查分词结果中是否有明显的句子边界
            for i, word in enumerate(sentences[:-1]):
                if word in ['。', '！', '？', '.', '!', '?', '；', ';']:
                    # 检查这个边界是否接近text_before和text_after的连接处
                    before_seg = ''.join(sentences[:i + 1])
                    if (len(before_seg) - len(text_before)) < 5:
                        return True
        else:
            # 英文文本，使用NLTK
            sentences = sent_tokenize(combined_text)
            # 检查是否能在组合文本中找到明确的句子分界点
            for sentence in sentences:
                if text_before.endswith(sentence) or text_after.startswith(sentence):
                    return True
    except:
        pass

    return False


def find_nearest_sentence_boundary(paragraphs_info, current_index, search_window=5):
    """寻找距离当前位置最近的句子边界"""
    best_index = -1
    min_distance = float('inf')

    # 向前查找
    for i in range(max(0, current_index - search_window), current_index + 1):
        if i > 0 and is_sentence_boundary(paragraphs_info[i - 1]['text'], paragraphs_info[i]['text']):
            distance = current_index - i
            if 0 <= distance < min_distance:
                min_distance = distance
                best_index = i

    # 向后查找
    for i in range(current_index + 1, min(len(paragraphs_info), current_index + search_window + 1)):
        if i > 0 and is_sentence_boundary(paragraphs_info[i - 1]['text'], paragraphs_info[i]['text']):
            distance = i - current_index
            if distance < min_distance:
                min_distance = distance
                best_index = i

    return best_index


def analyze_document(doc, debug_mode=False):
    """分析文档，提取段落信息和语义块"""
    paragraphs_info = []
    total_text = ""

    # 提取段落信息
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:  # 跳过空段落，但记录它们以保持索引一致
            paragraphs_info.append({
                'index': i,
                'text': "",
                'length': 0,
                'is_heading': False,
                'is_list_item': False,
                'ends_with_period': False,
                'length_category': "empty"
            })
            continue

        # 判断段落类型
        is_heading = paragraph.style.name.startswith(('Heading', '标题'))
        is_list_item = text.startswith(('•', '-', '*', '1.', '2.', '3.')) or (
                len(text) > 2 and text[0].isdigit() and text[1] == '.')

        # 段落长度分类
        length_category = "short" if len(text) < 50 else "medium" if len(text) < 200 else "long"

        # 尝试判断段落是否为完整句子结束
        ends_with_period = text.endswith(('。', '！', '？', '.', '!', '?', '；', ';'))

        # 将段落信息保存起来
        paragraphs_info.append({
            'index': i,
            'text': text,
            'length': len(text),
            'is_heading': is_heading,
            'is_list_item': is_list_item,
            'ends_with_period': ends_with_period,
            'length_category': length_category
        })

        total_text += text + " "

    # 使用NLTK尝试进行句子拆分（用于辅助判断语义边界）
    try:
        analyze_sentence_structure(total_text)
    except:
        pass

    # 预处理阶段：尝试将连续的短段落组合成语义块
    semantic_blocks = identify_semantic_blocks(paragraphs_info)

    return paragraphs_info, semantic_blocks


def analyze_sentence_structure(text):
    """分析文本的句子结构"""
    # 判断文本主要是中文还是英文
    if any(u'\u4e00' <= char <= u'\u9fff' for char in text):
        # 中文文本使用jieba分词和分句
        words = list(jieba.cut(text))
        sentences = []
        temp = ""
        for word in words:
            temp += word
            if word in ['。', '！', '？', '.', '!', '?']:
                sentences.append(temp)
                temp = ""
        if temp:  # 添加最后一个可能不完整的句子
            sentences.append(temp)
    else:
        # 英文文本使用NLTK
        sentences = sent_tokenize(text)

    return sentences


def identify_semantic_blocks(paragraphs_info):
    """识别文档中的语义块"""
    semantic_blocks = []
    temp_block = {'text': '', 'paragraphs': []}

    for i, para_info in enumerate(paragraphs_info):
        if para_info['length'] == 0:  # 跳过空段落
            continue

        # 如果是标题或列表项开始，结束前一个块
        if para_info['is_heading'] or (
                para_info['is_list_item'] and (i == 0 or not paragraphs_info[i - 1]['is_list_item'])):
            if temp_block['text']:
                semantic_blocks.append(temp_block)
                temp_block = {'text': '', 'paragraphs': []}

        # 添加到当前块
        temp_block['text'] += para_info['text'] + ' '
        temp_block['paragraphs'].append(i)

        # 如果段落结束一个完整句子，且不是短段落，也考虑结束当前块
        if para_info['ends_with_period'] and para_info['length'] > 100:
            semantic_blocks.append(temp_block)
            temp_block = {'text': '', 'paragraphs': []}

    # 添加最后一个块
    if temp_block['text']:
        semantic_blocks.append(temp_block)

    return semantic_blocks



def extract_elements_info(doc, table_length_factor=1.0, debug_mode=False):
    """
    按文档布局顺序返回 elements_info 列表，每一项：
       {
         'type'        : 'para' | 'table',
         'i_para'      : 段落索引 (仅 para 有),
         'i_table'     : 表格索引 (仅 table 有),
         'length'      : 文本字符数 * table_length_factor,
         'text'        : 纯文本(表格=单元格文本拼接),
         'is_heading'  : ...
         'is_list_item': ...
         'ends_with_period': ...
       }
    """
    elements = []
    para_idx = -1
    tbl_idx  = -1

    paragraph_map = {p._element: p for p in doc.paragraphs}
    table_map     = {t._element: t for t in doc.tables}

    for el in doc._element.body:
        # ---------- 段落 ----------
        if isinstance(el, CT_P):
            para_idx += 1
            p = paragraph_map[el]
            text = p.text.strip()
            is_heading = p.style.name.startswith(('Heading', '标题'))
            is_list_item = text.startswith(('•', '-', '*', '1.', '2.', '3.')) or (
                           len(text) > 2 and text[0].isdigit() and text[1] == '.')
            ends_with_period = text.endswith(('。', '！', '？', '.', '!', '?', '；', ';'))

            elements.append({
                'type': 'para',
                'i_para': para_idx,
                'i_table': None,
                'text': text,
                'length': len(text),
                'is_heading': is_heading,
                'is_list_item': is_list_item,
                'ends_with_period': ends_with_period
            })

        # ---------- 表格 ----------
        elif isinstance(el, CT_Tbl):
            tbl_idx += 1
            tbl = table_map[el]
            # 统计表格文本
            texts = []
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        texts.append(cell.text.strip())
            tbl_text = ' '.join(texts)
            tbl_len  = int(len(tbl_text) * table_length_factor)

            elements.append({
                'type': 'table',
                'i_para': None,
                'i_table': tbl_idx,
                'text': tbl_text,
                'length': tbl_len,
                'is_heading': False,
                'is_list_item': False,
                'ends_with_period': True      # 视为天然边界
            })

    if debug_mode:
        print(f"[extract] 生成 elements_info 共 {len(elements)} 条，其中表格 {tbl_idx+1} 个")

    return elements
